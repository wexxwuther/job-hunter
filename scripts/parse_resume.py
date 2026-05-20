#!/usr/bin/env python3
"""
parse_resume.py

Extract plain text from a resume in DOCX, PDF, MD, TXT, or HTML format. Lets the job-hunter
skill read resumes without depending on sibling docx/pdf skills (standalone contract).

For DOCX: uses python-docx (PEP 723 inline metadata declares the dependency).
For PDF: uses pypdf (also PEP 723).
For MD/TXT/HTML: reads as UTF-8 text. HTML is stripped of tags via stdlib html.parser.

Usage:
    python parse_resume.py resume.docx
    python parse_resume.py resume.pdf --out resume.txt
    python parse_resume.py resume.md --json

Run with `uv run` to get the inline-declared deps without polluting global Python:
    uv run scripts/parse_resume.py resume.docx
"""
# /// script
# requires-python = ">=3.10"
# dependencies = [
#   "python-docx>=1.1.0",
#   "pypdf>=4.0.0",
# ]
# ///

from __future__ import annotations

import argparse
import json
import re
import sys
from html.parser import HTMLParser
from pathlib import Path


class _HTMLStripper(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._chunks: list[str] = []
        self._skip = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style"}:
            self._skip += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style"} and self._skip > 0:
            self._skip -= 1

    def handle_data(self, data: str) -> None:
        if self._skip == 0:
            self._chunks.append(data)

    def text(self) -> str:
        return "".join(self._chunks)


def _strip_html(html: str) -> str:
    stripper = _HTMLStripper()
    stripper.feed(html)
    text = stripper.text()
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _parse_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "python-docx is not installed. Run with `uv run scripts/parse_resume.py ...` "
            "or `pip install python-docx`."
        ) from e
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError(
            "pypdf is not installed. Run with `uv run scripts/parse_resume.py ...` "
            "or `pip install pypdf`."
        ) from e
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _parse_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace").strip()


def _parse_html(path: Path) -> str:
    return _strip_html(path.read_text(encoding="utf-8", errors="replace"))


PARSERS = {
    ".docx": _parse_docx,
    ".pdf": _parse_pdf,
    ".md": _parse_text,
    ".markdown": _parse_text,
    ".txt": _parse_text,
    ".text": _parse_text,
    ".rst": _parse_text,
    ".html": _parse_html,
    ".htm": _parse_html,
}


def parse(path: str | Path) -> str:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"resume not found: {p}")
    if not p.is_file():
        raise IsADirectoryError(f"not a file: {p}")
    ext = p.suffix.lower()
    parser = PARSERS.get(ext)
    if not parser:
        raise ValueError(
            f"unsupported resume format: {ext!r}. Supported: {sorted(PARSERS.keys())}"
        )
    return parser(p)


def _basic_sections(text: str) -> dict[str, str]:
    section_patterns = [
        ("summary", r"(?:summary|profile|objective)"),
        ("experience", r"(?:experience|employment|work\s+history|professional\s+experience)"),
        ("education", r"education"),
        ("skills", r"(?:skills|technical\s+skills|core\s+competencies)"),
        ("certifications", r"(?:certifications?|licenses?)"),
        ("projects", r"projects?"),
    ]
    out: dict[str, str] = {}
    text_lines = text.splitlines()
    indices: list[tuple[str, int]] = []
    for i, line in enumerate(text_lines):
        stripped = line.strip()
        if not stripped or len(stripped) > 60:
            continue
        for key, pattern in section_patterns:
            if re.match(rf"^\s*{pattern}\s*$", stripped, re.IGNORECASE):
                indices.append((key, i))
                break
    indices.sort(key=lambda x: x[1])
    for j, (key, start) in enumerate(indices):
        end = indices[j + 1][1] if j + 1 < len(indices) else len(text_lines)
        body = "\n".join(text_lines[start + 1:end]).strip()
        if body:
            out[key] = body
    return out


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("resume", help="path to resume file")
    parser.add_argument("--out", default=None, help="write extracted text here (default stdout)")
    parser.add_argument("--json", action="store_true",
                        help="emit JSON with raw text plus best-effort section breakdown")
    args = parser.parse_args(argv)

    try:
        text = parse(args.resume)
    except (FileNotFoundError, IsADirectoryError, ValueError, RuntimeError) as e:
        print(f"error: {e}", file=sys.stderr)
        return 2

    if args.json:
        payload = {
            "source": str(Path(args.resume).resolve()),
            "format": Path(args.resume).suffix.lower(),
            "text": text,
            "sections": _basic_sections(text),
            "char_count": len(text),
            "line_count": text.count("\n") + 1,
        }
        out = json.dumps(payload, indent=2)
    else:
        out = text

    if args.out:
        Path(args.out).write_text(out + "\n", encoding="utf-8")
        print(f"wrote {len(out)} chars to {args.out}", file=sys.stderr)
    else:
        print(out)
    return 0


if __name__ == "__main__":
    sys.exit(main())
